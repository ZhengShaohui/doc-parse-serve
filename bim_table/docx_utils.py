#!/usr/bin/env python3
import os, sys, json
from pprint import pprint
import docx
from docx.shared import Inches, RGBColor


def _write_picture(document, filepath):
    document.add_picture(filepath, width=Inches(5))

def _write_text(document, text):
    p = document.add_paragraph(text)

def _write_table(document, table):
    nrow, ncol = table['shape']
    tb = document.add_table(rows=nrow, cols=ncol, style='Table Grid')
    tb.style.font.name = '宋体'
    # 单元格合并、赋值
    for _cell in table['cells']:
        cur_cells = []
        row, col = _cell['index']
        for i in range(_cell['cell_span'][0]):
            for j in range(_cell['cell_span'][1]):
                cur_cells.append([row+i-1, col+j-1])
        if len(cur_cells) > 1:
            i, j = cur_cells[0]
            cell_0 = tb.cell(i, j)
            _i,_j = cur_cells[-1]
            cell_x = tb.cell(_i, _j)
            cell_0.merge(cell_x)
        tb.cell(row-1, col-1).text = _cell['content']


def gen_docx(data, filename):
    document = docx.Document()
    for item in data:
        picture_file = item['filename']
        name = picture_file.split('/')[-1]
        document.add_heading('Title: '+name)
        h = document.add_heading('1-原图', 1)
        h.style.font.color.rgb = RGBColor(255, 0, 0)
        _write_picture(document, picture_file)
        h = document.add_heading('2-OCR结果', 1)
        h.style.font.color.rgb = RGBColor(255, 0, 0)
        objects = item.get('data', [])
        for _obj in objects:
            if _obj['type'] == 'text':
                _write_text(document, _obj['content'])
            elif _obj['type'] == 'table':
                _write_table(document, _obj)
                _write_text(document, '\n')
        _write_text(document, '\n')
    if not filename.endswith('.docx'):
        filename += '.docx'
    document.save(filename)




############################################
from docx.document import Document
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P
from docx.table import _Cell, Table
from docx.text.paragraph import Paragraph

def __iter_block_items(parent):
    if isinstance(parent, Document):
        parent_elm = parent.element.body
    elif isinstance(parent, _Cell):
        parent_elm = parent._tc
    else:
        raise ValueError("something's not right")

    for child in parent_elm.iterchildren():
        if isinstance(child, CT_P):
            yield Paragraph(child, parent)
        elif isinstance(child, CT_Tbl):
            yield Table(child, parent)

def __parse_table(table):
    """
    [table]:
    shape, type, cells, full_cells
    [cell]:
    'span', 'index', 'cell_span', 'content'
    """
    result = {}
    result['shape'] = (len(table.rows), len(table.columns))
    result['type'] = 'table'
    result['cells'] = []
    result['full_cells'] = []

    row_count, column_count = len(table.rows), len(table.columns)
    Cells = [[0 for _ in range(column_count)] for _ in range(row_count)]
    Cells_text = [['' for _ in range(column_count)] for _ in range(row_count)]
    _cells = table._cells
    for i in range(row_count):
        for j in range(column_count):
            cell_idx = j + (i * column_count)
            Cells[i][j] = _cells[cell_idx]
            Cells_text[i][j] = Cells[i][j].text + ' [' + str(i) + '-' + str(j) + ']'

    for i in range(row_count):
        for j in range(column_count):
            _cell = {}
            _cell['index'] = (i+1, j+1)
            _cell['cell_span'] = [1, 1]
            _cell['content'] = Cells_text[i][j]
            result['full_cells'].append(_cell)

    for i in range(row_count):
        for j in range(column_count):
            if Cells[i][j]._tc is Cells[i-1][j]._tc or Cells[i][j]._tc is Cells[i][j-1]._tc:
                pass
            else:
                _cell = {}
                _cell['index'] = (i+1, j+1)
                _cell['cell_span'] = [1, 1]
                _cell['content'] = Cells_text[i][j]
                result['cells'].append(_cell)

    #修正cell_span(对于合并单元格的)
    for _cell in result['cells']:
        i, j = _cell['index'][0]-1, _cell['index'][1]-1
        for _row in range(i+1, row_count):
            if Cells[i][j]._tc is Cells[_row][j]._tc:
                _cell['cell_span'][0] += 1
            else:
                break
        for _col in range(j+1, column_count):
            if Cells[i][j]._tc is Cells[i][_col]._tc:
                _cell['cell_span'][1] += 1
            else:
                break
    return result

def __split_items(data, filename):
    dirname = 'table_images/' + filename.split('/')[-1].replace('.docx', '')
    new_data = []
    cur_item = {}
    for x in data:
        flag = False
        if x['type'] == 'text' and x['content'].startswith('Title:'):
            if cur_item:
                new_data.append(cur_item)
            cur_item = {'filename':os.path.join(dirname, x['content'][6:].strip())}
        elif x['type'] == 'text' and x['content'] in ['1-原图', '2-OCR结果']:
            continue
        else:
            if 'data' in cur_item:
                cur_item['data'].append(x)
            else:
                cur_item['data'] = [x]
            pass
    if cur_item:
        new_data.append(cur_item)
    return new_data

def load_docx(filename):
    doc = docx.Document(filename)
    data = []
    for block in __iter_block_items(doc):
        if isinstance(block, Paragraph):
            text = block.text.strip()
            if text:
                data.append({'content':text, 'type':'text'})
        elif isinstance(block, Table):
            data.append(__parse_table(block))
    data = __split_items(data, filename)
    return data


if __name__ == '__main__':
    pass





